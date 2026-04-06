"""
generate_doc.py
---------------
Creates a comprehensive Hindi Word document explaining the
Payroll Reconciliation Tool codebase for a non-technical reader
with basic Python knowledge.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1.1)
section.right_margin  = Inches(1.1)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

# ── Colour palette ──────────────────────────────────────────────────────────
C_NAVY   = RGBColor(0x0D, 0x1B, 0x2E)
C_BLUE   = RGBColor(0x4F, 0x6B, 0xF5)
C_GREEN  = RGBColor(0x05, 0x96, 0x69)
C_AMBER  = RGBColor(0xD9, 0x77, 0x06)
C_RED    = RGBColor(0xDC, 0x26, 0x26)
C_GREY   = RGBColor(0x3D, 0x4C, 0x65)
C_LIGHT  = RGBColor(0xF0, 0xF3, 0xF8)
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
C_PURPLE = RGBColor(0x7C, 0x3A, 0xED)

# ── Helper functions ─────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'),   val.get('val', 'single'))
            el.set(qn('w:sz'),    val.get('sz', '4'))
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def add_heading(text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = color or C_NAVY
        if level == 1:
            run.font.size = Pt(18)
            run.bold = True
        elif level == 2:
            run.font.size = Pt(14)
            run.bold = True
        elif level == 3:
            run.font.size = Pt(12)
            run.bold = True
    return p

def add_para(text, bold=False, italic=False, color=None, size=11, indent=0, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color or C_GREY
    run.font.name = 'Mangal'
    return p

def add_bullet(text, level=0, color=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.8)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = color or C_GREY
    run.font.name = 'Mangal'
    return p

def add_numbered(text, color=None):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = color or C_GREY
    run.font.name = 'Mangal'
    return p

def add_code_box(lines, caption=None):
    """Light grey shaded paragraph for code/file references."""
    if caption:
        p = doc.add_paragraph()
        r = p.add_run(f'  📄 {caption}')
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = C_BLUE
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        r = p.add_run(line)
        r.font.name = 'Courier New'
        r.font.size = Pt(9.5)
        r.font.color.rgb = C_NAVY

def add_colored_box(text, bg_hex, text_color=None):
    """Single-cell table used as a coloured callout box."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg_hex)
    cell.width = Inches(6.3)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.4)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = text_color or C_NAVY
    run.font.name = 'Mangal'
    doc.add_paragraph()   # spacing after box

def add_flow_table(steps):
    """steps = list of (icon, title, detail) tuples — renders as a flow table."""
    tbl = doc.add_table(rows=len(steps), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = 'Table Grid'
    for i, (icon, title, detail) in enumerate(steps):
        # icon cell
        cell_icon = tbl.cell(i, 0)
        cell_icon.width = Cm(1.2)
        set_cell_bg(cell_icon, 'EEF2FF')
        p = cell_icon.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(icon)
        r.font.size = Pt(16)
        # content cell
        cell_txt = tbl.cell(i, 1)
        set_cell_bg(cell_txt, 'FFFFFF')
        p2 = cell_txt.paragraphs[0]
        r2 = p2.add_run(f'{title}\n')
        r2.bold = True
        r2.font.size  = Pt(11)
        r2.font.color.rgb = C_NAVY
        r2.font.name  = 'Mangal'
        r3 = p2.add_run(detail)
        r3.font.size  = Pt(10)
        r3.font.color.rgb = C_GREY
        r3.font.name  = 'Mangal'
    doc.add_paragraph()

def add_two_col_table(headers, rows, h_bg='1E3A5F', h_fg='FFFFFF'):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    # header row
    for j, hdr in enumerate(headers):
        cell = tbl.cell(0, j)
        set_cell_bg(cell, h_bg)
        p = cell.paragraphs[0]
        r = p.add_run(hdr)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(
            int(h_fg[0:2], 16), int(h_fg[2:4], 16), int(h_fg[4:6], 16))
        r.font.name = 'Mangal'
    # data rows
    for i, row in enumerate(rows):
        bg = 'F8FAFC' if i % 2 == 0 else 'FFFFFF'
        for j, val in enumerate(row):
            cell = tbl.cell(i + 1, j)
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(10)
            r.font.color.rgb = C_GREY
            r.font.name = 'Mangal'
    doc.add_paragraph()

def hr():
    """Horizontal divider paragraph."""
    p = doc.add_paragraph('─' * 80)
    p.runs[0].font.size  = Pt(7)
    p.runs[0].font.color.rgb = RGBColor(0xC8, 0xD3, 0xE3)
    doc.add_paragraph()

def page_break():
    doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('💼')
r.font.size = Pt(48)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Payroll Reconciliation Tool')
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = C_NAVY
r.font.name = 'Mangal'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('सम्पूर्ण कोड की व्याख्या — हिन्दी में')
r.font.size = Pt(18)
r.font.color.rgb = C_BLUE
r.font.name = 'Mangal'

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('(नये developer या non-technical व्यक्ति के लिए)')
r.font.size = Pt(13)
r.italic = True
r.font.color.rgb = C_GREY
r.font.name = 'Mangal'

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f'दस्तावेज़ तैयार: {datetime.date.today().strftime("%d %B %Y")}')
r.font.size = Pt(11)
r.font.color.rgb = C_GREY
r.font.name = 'Mangal'

page_break()

# ════════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS
# ════════════════════════════════════════════════════════════════════════════
add_heading('विषय-सूची (Table of Contents)', 1, C_NAVY)
doc.add_paragraph()

toc_items = [
    ('1.', 'इस Tool का उद्देश्य क्या है?', '3'),
    ('2.', 'Project की फाइल संरचना (Folder Structure)', '4'),
    ('3.', 'Tool कैसे काम करता है — पूरा Flow', '5'),
    ('4.', 'Frontend — User Interface की फाइलें', '6'),
    ('   4.1', 'index.html — मुख्य पन्ना', '6'),
    ('   4.2', 'style.css — रंग और डिज़ाइन', '7'),
    ('   4.3', 'app.js — मुख्य JS कोड', '8'),
    ('   4.4', 'upload.js — फाइल Upload का काम', '9'),
    ('   4.5', 'config.js — Configuration का काम', '10'),
    ('   4.6', 'results.js — Results दिखाने का काम', '11'),
    ('   4.7', 'history.js — History Page', '12'),
    ('5.', 'Backend — Server की Python फाइलें', '13'),
    ('   5.1', 'run.py — Server शुरू करना', '13'),
    ('   5.2', 'main.py — API का दरवाज़ा', '13'),
    ('   5.3', 'state.py — Session की याददाश्त', '14'),
    ('   5.4', 'db.py — MongoDB से जुड़ाव', '15'),
    ('   5.5', 'upload.py — फाइल प्राप्त करना', '16'),
    ('   5.6', 'columns.py — Column Mapping', '17'),
    ('   5.7', 'reconcile.py — Reconciliation चलाना', '18'),
    ('   5.8', 'mapping_config.py — Configuration Save करना', '19'),
    ('6.', 'Column Identifier — AI से Column पहचानना', '20'),
    ('7.', 'Processors — असली गणित का काम', '22'),
    ('8.', 'Excel Export — रिपोर्ट बनाना', '24'),
    ('9.', 'MongoDB — History कहाँ Save होती है?', '25'),
    ('10.', 'सभी फाइलों का आपस में संबंध (Connection Map)', '26'),
    ('11.', 'Dead Code — वो Code जो काम नहीं आ रहा', '27'),
    ('12.', 'सरल भाषा में पूरी प्रक्रिया (Summary)', '28'),
]
for num, title, pg in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    r1 = p.add_run(f'{num}  {title}')
    r1.font.size = Pt(11)
    r1.font.name = 'Mangal'
    r1.font.color.rgb = C_NAVY if not num.startswith('   ') else C_GREY

page_break()

# ════════════════════════════════════════════════════════════════════════════
#  SECTION 1 — PURPOSE
# ════════════════════════════════════════════════════════════════════════════
add_heading('1. इस Tool का उद्देश्य क्या है?', 1, C_NAVY)

add_para(
    'Payroll Reconciliation Tool एक ऐसा software है जो दो financial reports को आपस में मिलाकर '
    'देखता है — GL Report (General Ledger) और Payroll Register। इसका मुख्य काम यह पता लगाना है '
    'कि दोनों reports के आँकड़े (numbers) आपस में मेल खाते हैं या नहीं।'
)

doc.add_paragraph()
add_colored_box(
    '🏦  सरल भाषा में:\n'
    'मान लीजिए एक कंपनी के Accounts department के पास दो रजिस्टर हैं —\n'
    '• GL Report: Accounting system से निकली हुई report जिसमें हर GL Code '
    '(जैसे 5000 = Salaries) का total amount लिखा है।\n'
    '• Payroll Register: HR/Payroll system से निकली हुई report जिसमें हर employee '
    'के Pay Codes (जैसे Wages, Taxes, Benefits) का detail है।\n\n'
    'यह Tool दोनों को compare करता है और बताता है:\n'
    '✓ Match — दोनों में आँकड़े बराबर हैं\n'
    '⚠ Variance — कहीं अंतर है, जाँचें\n'
    '↓ Excel Report — पूरी detailed report download करें',
    'EEF2FF',
    C_NAVY
)

add_heading('इस Tool में क्या-क्या है?', 2, C_BLUE)

steps = [
    ('📂', 'Step 1 — Upload Files', 'User दोनों files (GL Report और Payroll Register) upload करता है। AI automatically columns पहचान लेता है।'),
    ('⚙',  'Step 2 — Configuration', 'User GL Codes को Payroll Pay Codes से map करता है (बताता है कि GL Code 5000 = Wages Pay Code से जुड़ा है)।'),
    ('▶',  'Step 3 — Run Reconciliation', 'Tool दोनों files को compare करता है और Variance निकालता है।'),
    ('📥', 'Step 4 — Download Report', 'Results को 6-sheet Excel file में export किया जा सकता है।'),
    ('⊙',  'Step 5 — History', 'हर run MongoDB में save हो जाता है। बाद में देख सकते हैं।'),
]
add_flow_table(steps)

page_break()

# ════════════════════════════════════════════════════════════════════════════
#  SECTION 2 — FOLDER STRUCTURE
# ════════════════════════════════════════════════════════════════════════════
add_heading('2. Project की फाइल संरचना (Folder Structure)', 1, C_NAVY)

add_para(
    'नीचे पूरे project की folder और file list है। हर एक का काम समझाया गया है।'
)
doc.add_paragraph()

add_code_box([
    'payroll_Reconciliation/',
    '│',
    '├── run.py                          ← Server शुरू करने की main file',
    '├── .env                            ← Secret keys (AWS, MongoDB) यहाँ होती हैं',
    '│',
    '├── config/',
    '│   ├── settings.py                 ← सभी settings एक जगह',
    '│   └── default_mapping.py          ← Default GL→Payroll mapping template',
    '│',
    '├── frontend/',
    '│   ├── index.html                  ← User जो screen देखता है (मुख्य पन्ना)',
    '│   └── static/',
    '│       ├── css/',
    '│       │   └── style.css           ← रंग, fonts, डिज़ाइन',
    '│       └── js/',
    '│           ├── app.js              ← मुख्य JS: navigation, session',
    '│           ├── upload.js           ← File upload page का logic',
    '│           ├── config.js           ← Configuration page का logic',
    '│           ├── results.js          ← Results page का logic',
    '│           └── history.js          ← History page का logic',
    '│',
    '├── backend/',
    '│   ├── api/',
    '│   │   ├── main.py                 ← FastAPI app, सभी routes register',
    '│   │   ├── state.py                ← Session data RAM में रखना',
    '│   │   ├── db.py                   ← MongoDB से connection',
    '│   │   └── routes/',
    '│   │       ├── upload.py           ← /api/upload endpoint',
    '│   │       ├── columns.py          ← /api/confirm-mapping endpoint',
    '│   │       ├── reconcile.py        ← /api/run, /api/download, /api/recon-history',
    '│   │       └── mapping_config.py   ← /api/mapping-config CRUD',
    '│   ├── column_identifier/',
    '│   │   ├── __init__.py             ← Column पहचान का orchestrator',
    '│   │   ├── fuzzy_matcher.py        ← String similarity से column match',
    '│   │   └── bedrock_identifier.py   ← AWS AI से column identify',
    '│   ├── processors/',
    '│   │   ├── gl_processor.py         ← GL Report को process करना',
    '│   │   ├── payroll_processor.py    ← Payroll Register process करना',
    '│   │   └── reconciliation_processor.py ← दोनों compare करना',
    '│   └── utils/',
    '│       ├── file_reader.py          ← Excel/CSV पढ़ना',
    '│       └── excel_exporter.py       ← Excel report बनाना',
    '│',
    '└── client_mappings/                ← Per-client config files यहाँ save होती हैं',
], caption='Project की पूरी File Structure')

page_break()

# ════════════════════════════════════════════════════════════════════════════
#  SECTION 3 — COMPLETE FLOW
# ════════════════════════════════════════════════════════════════════════════
add_heading('3. Tool कैसे काम करता है — पूरा Flow', 1, C_NAVY)

add_para(
    'जब कोई user Tool खोलता है तो कुछ steps में काम होता है। नीचे पूरी कहानी है:'
)
doc.add_paragraph()

flow_steps = [
    ('🌐', 'Browser Tool खोलता है',
     'User browser में http://localhost:8000 खोलता है। '
     'FastAPI server index.html file browser को भेजता है।'),
    ('🔑', 'Session बनती है',
     'Browser automatically /api/session को call करता है। '
     'Server एक unique ID (जैसे "abc-123-xyz") देता है। '
     'यह ID पूरे session में काम आती है — हर request में साथ जाती है।'),
    ('📂', 'GL File Upload होती है',
     'User GL Report file drag-drop करता है। '
     'upload.js file को server पर /api/upload/gl_report को भेजता है। '
     'Server file पढ़ता है, AI से columns पहचानता है, preview भेजता है।'),
    ('📂', 'Payroll File Upload होती है',
     'Same process Payroll Register के लिए। '
     'Columns identify होते हैं और user screen पर देखता है।'),
    ('✓', 'Column Mapping Confirm होती है',
     'User देखता है कि कौन सा column कौन सी role में है। '
     'Confirm button दबाने पर /api/confirm-mapping call होता है। '
     'Server mapping RAM में save कर लेता है।'),
    ('⚙', 'Configuration Set होती है',
     'Config page पर GL Code को Pay Code से link किया जाता है। '
     'Save करने पर /api/mapping-config call होता है — MongoDB या file में save होता है।'),
    ('▶', 'Reconciliation चलती है',
     'Run button दबाने पर /api/run call होता है। '
     'Server: GL process करता है → PR process करता है → Compare करता है → '
     'Excel बनाता है → Results browser को भेजता है → MongoDB में history save करता है।'),
    ('📥', 'Report Download होती है',
     'Download button से /api/download call होता है। '
     'Server Excel file stream करता है — browser download करता है।'),
    ('⊙', 'History में देखना',
     'History page खोलने पर /api/recon-history call होता है। '
     'MongoDB से सभी पुराने runs की list आती है। '
     'किसी पर click करने पर /api/recon-history/{id} से पूरा result load होता है।'),
]
add_flow_table(flow_steps)

page_break()

# ════════════════════════════════════════════════════════════════════════════
#  SECTION 4 — FRONTEND FILES
# ════════════════════════════════════════════════════════════════════════════
add_heading('4. Frontend — User Interface की फाइलें', 1, C_NAVY)

add_para(
    '"Frontend" का मतलब वो सब कुछ जो user browser में देखता है — '
    'रंग, buttons, tables, inputs। ये files browser में run होती हैं, server पर नहीं।'
)
doc.add_paragraph()

# 4.1 index.html
add_heading('4.1  index.html — मुख्य पन्ना', 2, C_BLUE)

add_para(
    'यह file एक Single Page Application (SPA) है — मतलब पूरा UI एक ही HTML file में है। '
    'Pages बदलते नहीं, बल्कि JavaScript show/hide करती है।'
)

add_two_col_table(
    ['Section का नाम', 'क्या होता है वहाँ'],
    [
        ['Landing Page', 'Tool खोलने पर पहला स्वागत screen। "Open Tool" button है।'],
        ['Sidebar / Navbar', 'बायीं तरफ का menu। Dashboard, Upload, Config, Results, History के links।'],
        ['Page 1 — Dashboard', '4 tiles: Config status, GL file status, PR file status, Results status।'],
        ['Page 2 — Upload Files', 'दो drag-drop zones — GL Report और Payroll Register के लिए।'],
        ['Page 3 — Configuration', 'GL Code → Pay Code mapping table। AI generator, save button।'],
        ['Page 4 — Results', 'Pre-flight checklist, Run button, Summary metrics, 5 tabs में results।'],
        ['Page 5 — History', 'MongoDB से आये पुराने runs की cards।'],
        ['Loading Overlay', 'हर API call के दौरान spinner दिखाई देता है।'],
    ]
)

add_colored_box(
    '💡 Important बात:\n'
    'index.html में कोई JavaScript logic नहीं लिखी है। '
    'सिर्फ HTML structure है। सभी logic अलग-अलग .js files में है।\n'
    'File के नीचे 5 script tags हैं:\n'
    '<script src="app.js"> + upload.js + config.js + results.js + history.js',
    'FFFBEB', C_AMBER
)
doc.add_paragraph()

# 4.2 style.css
add_heading('4.2  style.css — रंग और डिज़ाइन', 2, C_BLUE)

add_para(
    'यह file पूरे tool का visual design define करती है। '
    'सबसे ऊपर CSS Variables (Design Tokens) हैं — एक जगह color बदलो, पूरा tool बदल जाता है।'
)

add_two_col_table(
    ['Variable', 'रंग का काम'],
    [
        ['--blue: #4F6BF5', 'Primary buttons, active states, GL headers'],
        ['--green: #059669', 'Success states, "Clean" badges, Confirm buttons'],
        ['--red: #DC2626', 'Errors, Variance alerts, "Missing" badges'],
        ['--amber: #D97706', 'Warnings, Pending states'],
        ['--surface: #FFFFFF', 'Cards और panels का background'],
        ['--bg: #F0F3F8', 'पूरे page का background'],
        ['--text-1: #0A0F1E', 'Headings — गहरा रंग'],
        ['--text-3: #7E93B0', 'Hints और labels — हल्का रंग'],
    ]
)

add_para('CSS Sections जो important हैं:', bold=True)
add_bullet('Sidebar CSS — बायाँ navigation panel, white background')
add_bullet('Card CSS — हर section का box (shadow, border-radius)')
add_bullet('Upload Zone CSS — drag-drop area, hover animation')
add_bullet('Config Table CSS — Ghost inputs (transparent जब hover न हो), section headers')
add_bullet('Result Table CSS — Green rows (match), Red rows (variance), Blue (total)')
add_bullet('History Cards CSS — Cards grid, stats tiles')
add_bullet('Buttons CSS — btn-blue, btn-ghost, btn-files, btn-success')
doc.add_paragraph()

# 4.3 app.js
add_heading('4.3  app.js — मुख्य JS Code (Core)', 2, C_BLUE)

add_para(
    'यह file सबसे पहले load होती है। इसमें वो सब है जो पूरे tool में हर जगह काम आता है।'
)

add_two_col_table(
    ['Function का नाम', 'क्या करता है'],
    [
        ['state = {...}', 'Global object — पूरे session का data यहाँ रहता है (sessionId, files, results, etc.)'],
        ['createSession()', 'Server को /api/session call करता है, session ID मिलती है'],
        ['navigate(page)', 'एक page से दूसरे page पर जाना (CSS show/hide से)'],
        ['updateDashboard()', 'Dashboard की 4 tiles को current state के हिसाब से update करता है'],
        ['setupNavigation()', 'Sidebar के nav items पर click listener लगाता है'],
        ['setupTabs()', 'Results page के tabs (Reconciliation, GL Pivot, etc.) का logic'],
        ['startNewReconciliation()', 'पूरा state reset — नई process शुरू करना'],
        ['checkDbStatus()', '/api/db-status call करके sidebar में MongoDB dot दिखाना'],
        ['showLoading(msg)', 'Processing spinner दिखाना'],
        ['hideLoading()', 'Spinner हटाना'],
        ['getClient()', 'Sidebar input से client name पढ़ना'],
        ['esc(str)', 'HTML injection से बचाव (security)'],
    ]
)

add_colored_box(
    '📌 state object — यह सबसे ज़रूरी है:\n'
    'state.sessionId        — Server से मिला unique ID\n'
    'state.files            — कौन सी file upload हुई\n'
    'state.confirmed        — क्या column mapping confirm हुई\n'
    'state.uploadData       — Upload के बाद server से आया data\n'
    'state.mappingRows      — Config table की rows\n'
    'state.glCodeTitles     — GL file से पढ़े गये GL Codes\n'
    'state.prCodeTypes      — PR file से पढ़े गये Pay Codes\n'
    'state.results          — Reconciliation का result\n'
    'state.configSaved      — क्या config save हुई?',
    'EEF2FF', C_NAVY
)
doc.add_paragraph()

# 4.4 upload.js
add_heading('4.4  upload.js — File Upload का Logic', 2, C_BLUE)

add_para(
    'यह file Upload page (Step 1) का पूरा काम करती है।'
)

add_two_col_table(
    ['Function', 'काम'],
    [
        ['setupUploadZones()', 'दोनों zones पर drag-drop और click listeners लगाता है'],
        ['handleFileUpload()', 'File मिलने पर server को भेजता है, response handle करता है'],
        ['showSheetSelector()', 'Excel में multiple sheets हों तो selector दिखाता है'],
        ['renderFileSuccess()', 'File upload होने के बाद zone में filename और row count दिखाता है'],
        ['renderPreview()', 'File के पहले 5 rows की preview table बनाता है'],
        ['renderColumnAssignment()', 'AI ने जो columns detect किये, वो table में दिखाता है। User change कर सकता है।'],
        ['onRoleSelectChange()', 'User कोई column change करे तो badge "Manual" हो जाता है'],
        ['confirmColumnMapping()', 'Confirm button: /api/confirm-mapping को call करता है'],
        ['resetUpload()', 'File हटाना — zone को fresh करता है'],
        ['updateStepStatus()', 'Card header में animated status dot update करता है (Waiting → Pending → ✓ Confirmed)'],
    ]
)

add_colored_box(
    '🔄 Upload का Flow (Detail में):\n'
    '1. User file drag करता है → handleFileUpload() call होती है\n'
    '2. FormData बनती है (file + session_id + client_name)\n'
    '3. POST /api/upload/gl_report को भेजा जाता है\n'
    '4. Server file पढ़ता है, AI columns detect करता है\n'
    '5. Response आने पर renderFileSuccess() + renderPreview() + renderColumnAssignment() call होती हैं\n'
    '6. User columns देखता है, Confirm दबाता है → /api/confirm-mapping\n'
    '7. Confirm होने पर fetchGLCodes() या fetchPRCodes() call होती है (config.js में)',
    'ECFDF5', C_GREEN
)
doc.add_paragraph()

# 4.5 config.js
add_heading('4.5  config.js — Configuration का Logic', 2, C_BLUE)

add_para(
    'यह file Configuration page (Step 2) का पूरा काम करती है — '
    'mapping table render करना, GL/PR codes load करना, save करना।'
)

add_two_col_table(
    ['Function', 'काम'],
    [
        ['fetchGLCodes()', '/api/gl-codes call — GL file से unique GL Codes लाना, state में save करना'],
        ['fetchPRCodes()', '/api/pr-codes call — PR file से unique Pay Codes लाना'],
        ['showFilesLoadedBanner()', 'Green banner दिखाना "X GL codes loaded from files"'],
        ['autoPopulateConfigFromFiles()', 'जब दोनों files confirm हों और config unsaved हो तो auto rows बनाना'],
        ['buildSuggestedRows()', 'GL Codes को ranges के हिसाब से steps में group करना (5000-5099 = Earning)'],
        ['loadFromFiles()', '"Load from Files" button — confirm dialog के बाद rows replace करना'],
        ['onGLCodeInput()', 'GL Code type करने पर auto-fill GL Title (datalist से)'],
        ['onPayCodeInput()', 'Pay Code type करने पर Code Type auto-select'],
        ['renderMappingConfigTable()', 'Mapping table HTML generate करना — group rows, data rows, datalists'],
        ['collectMappingRows()', 'Table से current data DOM से पढ़ना (save से पहले)'],
        ['saveMappingConfig()', '/api/mapping-config POST — MongoDB या file में save'],
        ['loadMappingConfig()', '/api/mapping-config GET — saved config load करना'],
        ['resetToDefault()', 'Config delete करके default template reload'],
        ['generateMappingWithAI()', 'AWS Bedrock AI से config generate करना'],
    ]
)

add_colored_box(
    '🤖 Auto-populate Logic (नया feature):\n'
    'जब दोनों files confirm हों, तो:\n'
    '1. fetchGLCodes() और fetchPRCodes() call होती हैं\n'
    '2. GL Codes server से आते हैं, state.glCodeTitles में save\n'
    '3. autoPopulateConfigFromFiles() check करता है:\n'
    '   • क्या config already saved है? → तो कुछ नहीं\n'
    '   • क्या user ने manually rows edit किये? → तो कुछ नहीं\n'
    '   • नहीं तो: buildSuggestedRows() → table में डाल दो\n'
    '4. GL Code 5000-5099 → "A. Earning / Gross Wages" step में\n'
    '   GL Code 5100-5199 → "B. Benefits" step में, etc.',
    'F5F3FF', C_PURPLE
)
doc.add_paragraph()

# 4.6 results.js
add_heading('4.6  results.js — Results दिखाने का Logic', 2, C_BLUE)

add_para(
    'यह file Results page (Step 3) का पूरा काम करती है।'
)

add_two_col_table(
    ['Function', 'काम'],
    [
        ['updatePreflight()', 'Run से पहले 3-item checklist update करना (GL ✓, PR ✓, Config ✓)'],
        ['runReconciliation()', '/api/run call — reconciliation शुरू करना, results receive करना'],
        ['renderResults()', 'सभी render functions को एक साथ call करना'],
        ['renderStatusBanner()', 'Run के बाद top पर बड़ा banner — Green (Clean) या Red (Variance)'],
        ['renderSummaryMetrics()', '4 metric tiles: Total Lines, Matched, Variances, Variance Amount (animated counter)'],
        ['renderReconGrouped()', 'Reconciliation table — section headers (blue), match rows (green), variance rows (red)'],
        ['renderCombinedPivot()', 'GL और PR का combined table — Blue GL side, Green PR side'],
        ['renderResultTable()', 'GL Pivot और PR Pivot tables'],
        ['renderUnmapped()', '"Unmapped" tab — वो GL codes जो config में नहीं हैं'],
        ['downloadExcel()', '/api/download call — Excel file download'],
        ['showDownloadBar()', 'Download card में client name और period दिखाना'],
    ]
)
doc.add_paragraph()

# 4.7 history.js
add_heading('4.7  history.js — History Page का Logic', 2, C_BLUE)

add_two_col_table(
    ['Function', 'काम'],
    [
        ['loadHistory()', '/api/recon-history call — सभी runs की list MongoDB से लाना'],
        ['filterHistory()', 'Search box से client name से filter करना'],
        ['renderHistoryList()', 'Runs को client groups में cards की form में render करना'],
        ['renderRunCard()', 'एक run का card HTML बनाना (period, date, stats, files, buttons)'],
        ['viewHistoryRecord()', '/api/recon-history/{id} call — पुराना पूरा result load करना'],
        ['startNewReconForClient()', 'History से "New Run" — client name pre-fill करके fresh start'],
    ]
)

page_break()

# ════════════════════════════════════════════════════════════════════════════
#  SECTION 5 — BACKEND FILES
# ════════════════════════════════════════════════════════════════════════════
add_heading('5. Backend — Server की Python Files', 1, C_NAVY)

add_para(
    '"Backend" वो Python code है जो server पर run होता है। '
    'Browser directly data process नहीं कर सकता — वो server को request भेजता है, '
    'server Python में process करके result वापस भेजता है।\n\n'
    'यह project FastAPI framework use करता है — यह Python का एक modern web framework है '
    'जो API endpoints बनाने में use होता है।'
)
doc.add_paragraph()

# 5.1 run.py
add_heading('5.1  run.py — Server शुरू करना', 2, C_BLUE)
add_code_box([
    '# run.py का काम — बस server start करना',
    'import uvicorn',
    'uvicorn.run("backend.api.main:app", host="0.0.0.0", port=8000)',
], caption='run.py')
add_para(
    'यह एकदम simple file है। इसे run करने पर server start होता है: python run.py\n'
    'uvicorn एक ASGI server है जो FastAPI को run करता है। Port 8000 पर browser connect करता है।'
)
doc.add_paragraph()

# 5.2 main.py
add_heading('5.2  main.py — API का दरवाज़ा (Gateway)', 2, C_BLUE)

add_para(
    'यह file FastAPI app बनाती है और सभी routes को register करती है। '
    'जब browser कोई URL request करता है, main.py उसे सही route पर भेजता है।'
)

add_two_col_table(
    ['Route / Endpoint', 'क्या होता है'],
    [
        ['GET  /', 'Browser को index.html भेजना (frontend serve करना)'],
        ['GET  /static/*', 'CSS, JS, images files serve करना'],
        ['POST /api/session', 'नई session बनाना, unique ID देना'],
        ['POST /api/session/{id}/reset', 'Session का data clear करना (New Reconciliation)'],
        ['POST /api/upload/{file_type}', 'File receive करना (upload.py में handle)'],
        ['POST /api/confirm-mapping', 'Column mapping save करना (columns.py में handle)'],
        ['POST /api/run', 'Reconciliation चलाना (reconcile.py में handle)'],
        ['GET  /api/download', 'Excel file stream करना'],
        ['GET  /api/mapping-config', 'Saved config load करना'],
        ['POST /api/mapping-config', 'Config save करना'],
        ['DELETE /api/mapping-config', 'Config delete करना'],
        ['GET  /api/gl-codes', 'GL file से unique codes निकालना'],
        ['GET  /api/pr-codes', 'PR file से unique pay codes निकालना'],
        ['GET  /api/recon-history', 'MongoDB से history list'],
        ['GET  /api/recon-history/{id}', 'एक specific run का full result'],
        ['GET  /api/db-status', 'MongoDB connected है या नहीं'],
    ]
)
doc.add_paragraph()

# 5.3 state.py
add_heading('5.3  state.py — Session की याददाश्त (In-Memory Store)', 2, C_BLUE)

add_para(
    'Python RAM में एक dictionary (_store) है जो हर session का data रखती है। '
    'Server restart होने पर यह data खो जाता है — इसीलिए MongoDB भी use होता है।'
)

add_code_box([
    '# state.py का structure',
    '_store = {',
    '  "abc-123-xyz": {           ← session_id',
    '    "files": {',
    '      "gl_report": {',
    '        "df": DataFrame,     ← Excel/CSV का पूरा data',
    '        "filename": "GL.xlsx",',
    '        "header_row": 2',
    '      },',
    '      "payroll_register": { ... }',
    '    },',
    '    "mappings": {',
    '      "gl_report": {',
    '        "AcctCode": "gl_code",  ← actual column → semantic role',
    '        "AcctName": "gl_title",',
    '        ...',
    '      }',
    '    },',
    '    "results": {',
    '      "excel_bytes": b"...",  ← download के लिए ready Excel',
    '      "summary_stats": {...},',
    '      ...',
    '    }',
    '  }',
    '}',
], caption='_store dictionary का structure')

add_two_col_table(
    ['Function', 'काम'],
    [
        ['new_session()', 'UUID generate करके नई empty session बनाना'],
        ['get(sid)', 'Session data return करना'],
        ['set_file(sid, type, df, filename)', 'Uploaded file data save करना'],
        ['set_mapping(sid, type, mapping)', 'Confirmed column mapping save करना'],
        ['set_results(sid, results)', 'Reconciliation results save करना'],
        ['get_file / get_mapping / get_results', 'Saved data retrieve करना'],
        ['reset_session(sid)', 'Files, mappings, results clear करना (session ID रहती है)'],
    ]
)
doc.add_paragraph()

# 5.4 db.py
add_heading('5.4  db.py — MongoDB से जुड़ाव', 2, C_BLUE)

add_para(
    'यह file MongoDB से connection manage करती है। '
    'अगर .env में MONGO_URI set नहीं है तो यह automatically file-based storage use करती है।'
)

add_code_box([
    '# .env में यह add करें MongoDB के लिए:',
    'MONGO_URI=mongodb://localhost:27017',
    'MONGO_DB=payroll_recon',
], caption='.env file में MongoDB setup')

add_para('MongoDB में दो collections हैं:', bold=True)
add_bullet('mapping_configs — हर client की GL→PR mapping configuration')
add_bullet('recon_history — हर reconciliation run का record (full results सहित)')

add_colored_box(
    '⚡ Lazy Connection:\n'
    'db.py पहली बार call होने पर connect करती है, हर request पर नहीं।\n'
    'Connection fail हो तो silently file storage पर fall back हो जाती है — '
    'tool बिना MongoDB के भी काम करता है।',
    'FFFBEB', C_AMBER
)
doc.add_paragraph()

# 5.5 upload.py
add_heading('5.5  backend/routes/upload.py — File Receive करना', 2, C_BLUE)

add_para(
    'जब browser file भेजता है, यह endpoint उसे receive करता है।'
)

add_flow_table([
    ('📥', 'File Receive होती है',
     'Multipart form data में file + session_id + client_name + sheet_name आता है।'),
    ('🔍', 'Format Check',
     'File extension .xlsx/.csv/.ods आदि होनी चाहिए।'),
    ('📖', 'File पढ़ना (file_reader.py)',
     'file_reader.py auto-detect करता है header row कहाँ है। '
     'DataFrame बनता है।'),
    ('🤖', 'AI Column Identification (column_identifier)',
     'identify_columns() call होती है — '
     'पहले fuzzy match, फिर जरूरत पड़े तो AWS Bedrock AI।'),
    ('💾', 'Session में Save',
     'state.set_file() से DataFrame और filename session में store।'),
    ('↩', 'Response Browser को',
     'columns list, auto-detected mapping, confidence scores, preview rows भेजे जाते हैं।'),
])
doc.add_paragraph()

# 5.6 columns.py
add_heading('5.6  backend/routes/columns.py — Column Mapping Confirm करना', 2, C_BLUE)

add_para(
    'जब user column assignment check करके Confirm button दबाता है, यह endpoint call होता है।'
)

add_code_box([
    '# confirm-mapping को यह data आता है:',
    '{',
    '  "session_id": "abc-123",',
    '  "file_type": "gl_report",',
    '  "mapping": {',
    '    "AcctCode": "gl_code",    ← actual col → role',
    '    "AcctName": "gl_title",',
    '    "TransSource": "trans_source",',
    '    "NetAmt": "net_amount"',
    '  },',
    '  "client_name": "ClientABC",',
    '  "save_cache": true',
    '}',
], caption='/api/confirm-mapping का request')

add_para(
    'Mapping session में save होती है। अगर save_cache true हो तो '
    'client_mappings/ folder में JSON file बन जाती है — '
    'अगली बार same file structure upload हो तो AI की जरूरत नहीं।'
)
doc.add_paragraph()

# 5.7 reconcile.py
add_heading('5.7  backend/routes/reconcile.py — Reconciliation का दिल', 2, C_BLUE)

add_para(
    'यह file सबसे important है। /api/run endpoint यहाँ है जो पूरी '
    'reconciliation process चलाता है।'
)

add_flow_table([
    ('✅', 'Validation',
     'Session exist करती है? GL और PR दोनों upload और confirm हैं?'),
    ('📋', 'Config Load करना',
     '_load_config(client_name) — MongoDB या file से GL→PR mapping rows load करना। '
     'build_lookups_from_config() से 3 dictionaries बनती हैं: gl_lookup, pr_lookup, gl_pr_amount।'),
    ('📊', 'GL Process करना',
     'process_gl(df, col_map, gl_lookup) — GL file को filter और pivot करना। '
     'PRS transactions select करना, GL Code से group करना।'),
    ('📊', 'PR Process करना',
     'process_payroll(df, col_map, pr_lookup) — Payroll Register को process करना। '
     'Pay codes को GL codes से map करके pivot बनाना।'),
    ('🔢', 'Reconciliation',
     'build_reconciliation() — GL pivot और PR pivot को compare करना। '
     'हर row पर Variance निकालना।'),
    ('📈', 'Summary Stats',
     'get_summary_stats() — Total lines, matched, variances, total variance, is_clean।'),
    ('💾', 'MongoDB History',
     'db["recon_history"].insert_one() — पूरा result MongoDB में save।'),
    ('📥', 'Excel बनाना',
     'export_to_excel() — 6-sheet workbook: GL Mapped, PR Mapped, GL Pivot, PR Pivot, '
     'Reconciliation, Payroll Process।'),
    ('↩', 'Response',
     'JSON में: summary_stats, recon_table, gl_pivot, pr_pivot, unmapped_gl, unmapped_pr।'),
])

add_colored_box(
    '🔑 gl_lookup, pr_lookup, gl_pr_amount — ये तीनों क्या हैं?\n\n'
    'gl_lookup = {\n'
    '  "5000": {"gl_title": "Salaries", "recon_step": "A. Earning", "code_type": "EARNING"}\n'
    '}\n'
    '→ GL Code से उसका step और type मिलता है।\n\n'
    'pr_lookup = {\n'
    '  ("Wages", "EARNING"): "5000 - Salaries"\n'
    '}\n'
    '→ Pay Code + Code Type से उसका GL Code मिलता है।\n\n'
    'gl_pr_amount = {\n'
    '  "2115": {"5000 - Salaries": "EETax", "5001 - OT": "EETax & ERTax"}\n'
    '}\n'
    '→ Same GL Code के लिए अलग-अलग amount columns use होती हैं।',
    'EEF2FF', C_NAVY
)
doc.add_paragraph()

# 5.8 mapping_config.py
add_heading('5.8  backend/routes/mapping_config.py — Config CRUD', 2, C_BLUE)

add_para(
    'यह file Configuration page की CRUD operations handle करती है।'
)

add_two_col_table(
    ['Endpoint', 'काम'],
    [
        ['GET /api/mapping-config', 'Client की saved config load करना (MongoDB → file → default)'],
        ['POST /api/mapping-config', 'Config save/update करना (MongoDB → file fallback)'],
        ['DELETE /api/mapping-config', 'Config delete करना, default template restore'],
        ['POST /api/generate-mapping', 'AWS Bedrock AI से config generate करना'],
        ['GET /api/mapping-config/template', 'Default 80-row template return करना'],
    ]
)

add_code_box([
    '# mapping row का structure (7 fields):',
    '{',
    '  "recon_step": "A. Earning / Gross Wages",',
    '  "gl_code": "5000",',
    '  "gl_title": "Salaries & Wages",',
    '  "pay_code": "Wages",',
    '  "pay_code_title": "Regular Wages",',
    '  "amount_column": "EarnAmt",',
    '  "code_type": "EARNING"',
    '}',
], caption='Mapping Row का Format')
doc.add_paragraph()

page_break()

# ════════════════════════════════════════════════════════════════════════════
#  SECTION 6 — COLUMN IDENTIFIER
# ════════════════════════════════════════════════════════════════════════════
add_heading('6. Column Identifier — AI से Column पहचानना', 1, C_NAVY)

add_para(
    'यह feature बहुत important है। अलग-अलग companies की files में columns के नाम अलग होते हैं:\n'
    '• किसी की GL file में column "AcctCode" है, किसी की "Account Number", किसी की "GL_Code"\n'
    '• किसी की PR file में "PayCode" है, किसी की "Pay Code", किसी की "p_code"\n\n'
    'Column Identifier automatically समझता है कि कौन सा column कौन सी role निभाता है।'
)

add_heading('3 Step Process:', 2, C_BLUE)

add_flow_table([
    ('1️⃣', 'Cache Check (सबसे पहले)',
     'क्या इस client ने पहले यह file structure upload की है? '
     'अगर हाँ तो cache से direct mapping लो — AI की जरूरत नहीं। '
     'Cache key = client_name + file_type + MD5(column names)'),
    ('2️⃣', 'Fuzzy Matching (rapidfuzz library)',
     'हर column का नाम known aliases से compare होता है (85% threshold)। '
     'जैसे "AcctCode" → 92% match with "gl_code" aliases → confirmed।\n'
     'Fuzzy matching word order भी handle करती है ("Code Account" = "Account Code")।'),
    ('3️⃣', 'AWS Bedrock AI (जब fuzzy fail हो)',
     'जो columns 85% से कम match हों, उनके लिए AI call होती है। '
     '10 sample rows भेजे जाते हैं Claude Haiku को। '
     'AI बताता है: {role: "gl_code", confidence: 0.95, reason: "..."}'),
])

add_two_col_table(
    ['File', 'काम'],
    [
        ['__init__.py', 'identify_columns() — orchestrator। तीनों steps को order में call करता है।'],
        ['fuzzy_matcher.py', 'rapidfuzz library से token_set_ratio calculate करना।'],
        ['bedrock_identifier.py', 'boto3 से AWS Bedrock (Claude Haiku) call करना।'],
        ['mapping_cache.py', 'Cache read/write — JSON files client_mappings/ folder में।'],
    ]
)

add_colored_box(
    '📚 Settings में बहुत सारे Column Aliases हैं:\n'
    'config/settings.py में हर semantic role के लिए 10-20 known column names हैं।\n'
    'जैसे "gl_code" role के aliases:\n'
    '"AcctCode", "Account Number", "GL Code", "GL_Code", "account_code", '
    '"acct_no", "AccountNo", "GLCode" — इनमें से कोई भी मिले तो पहचान लेगा।',
    'F0F3F8', C_NAVY
)
doc.add_paragraph()

page_break()

# ════════════════════════════════════════════════════════════════════════════
#  SECTION 7 — PROCESSORS
# ════════════════════════════════════════════════════════════════════════════
add_heading('7. Processors — असली गणित का काम', 1, C_NAVY)

add_para(
    'यहाँ असली data processing होती है — pandas DataFrames को transform करना।'
)

# GL Processor
add_heading('7.1  gl_processor.py — GL Report Process करना', 2, C_BLUE)

add_para('यह file GL Report DataFrame को 3 outputs में बदलती है:')
add_bullet('gl_mapped — Filter और clean हुई GL rows')
add_bullet('gl_pivot — GL Code से group करके summed amounts')
add_bullet('unmapped_gl — वो GL codes जो config में नहीं हैं')

add_code_box([
    '# GL Processor का Main Logic',
    '',
    '# Step 1: PRS transactions filter करना',
    '# TransSource column में "PRS" होना चाहिए',
    'df = df[df[trans_source_col].str.contains("PRS", na=False)]',
    '',
    '# Step 2: GL Code से group और sum',
    'gl_pivot = df.groupby(gl_code_col)[net_amount_col].sum().reset_index()',
    '',
    '# Step 3: gl_lookup से step और title add करना',
    'gl_pivot["Reconciliation Mapping"] = gl_pivot["GL Code"].map(',
    '    lambda c: gl_lookup.get(c, {}).get("recon_step", "Unmapped")',
    ')',
], caption='gl_processor.py का simplified logic')
doc.add_paragraph()

# PR Processor
add_heading('7.2  payroll_processor.py — Payroll Register Process करना', 2, C_BLUE)

add_para('Payroll Register में बहुत सारी rows होती हैं (हर employee, हर pay code)। इन्हें summarize करना होता है।')

add_code_box([
    '# PR Processor का Main Logic',
    '',
    '# Step 1: हर (pay_code, code_type) combination के लिए',
    '#          amounts sum करना',
    'pr_grouped = df.groupby([pay_code_col, code_type_col]).agg({',
    '    earn_col: "sum",',
    '    bene_col: "sum",',
    '    deduc_col: "sum",',
    '    eetax_col: "sum",',
    '    ertax_col: "sum"',
    '}).reset_index()',
    '',
    '# Step 2: pr_lookup से GL Code find करना',
    'pr_grouped["Reconciliation Mapping"] = pr_grouped.apply(',
    '    lambda r: pr_lookup.get((r[pay_code_col], r[code_type_col]), "Unmapped"),',
    '    axis=1',
    ')',
], caption='payroll_processor.py का simplified logic')
doc.add_paragraph()

# Reconciliation Processor
add_heading('7.3  reconciliation_processor.py — Compare और Variance निकालना', 2, C_BLUE)

add_para(
    'यह सबसे important processor है। GL Pivot और PR Pivot को compare करता है।'
)

add_two_col_table(
    ['GL Code Range', 'Account Type', 'Variance Formula'],
    [
        ['1xxx (1000-1999)', 'Bank/Asset accounts', 'Variance = GL Amount − PR Amount'],
        ['2xxx (2000-2999)', 'Liability accounts', 'Variance = GL Amount + PR Amount'],
        ['5xxx/6xxx (5000+)', 'Expense accounts', 'Variance = GL Amount − PR Amount'],
    ]
)

add_colored_box(
    '🔢 Sign Convention क्यों अलग है?\n'
    'Accounting में Liabilities (2xxx) credit side पर होती हैं — इसलिए sign opposite होता है।\n'
    'जैसे GL code 2115 (Federal Tax Payable) GL में negative होगा (credit), '
    'PR में positive — इसलिए GL + PR = 0 होना चाहिए।',
    'FEF2F2', C_RED
)

add_para('gl_pr_amount Dictionary का Special Use:', bold=True)
add_para(
    'एक GL code के लिए multiple amount columns use हो सकती हैं:\n'
    '• GL 2115 — FIT pay code → सिर्फ EETax column\n'
    '• GL 2115 — SS pay code → EETax + ERTax दोनों columns\n'
    'gl_pr_amount dictionary यही बताती है कि कब कौन सा column use करना है।'
)
doc.add_paragraph()

page_break()

# ════════════════════════════════════════════════════════════════════════════
#  SECTION 8 — EXCEL EXPORT
# ════════════════════════════════════════════════════════════════════════════
add_heading('8. Excel Export — Report बनाना (excel_exporter.py)', 1, C_NAVY)

add_para(
    'Reconciliation के बाद एक professional 6-sheet Excel workbook बनती है।'
)

add_two_col_table(
    ['Sheet नाम', 'क्या है उसमें'],
    [
        ['1. GL Mapped', 'Filter और mapped GL transactions — सभी columns'],
        ['2. PR Mapped', 'Filter और mapped Payroll transactions'],
        ['3. GL Pivot', 'GL Code से grouped summary — Sum of Net Amount'],
        ['4. PR Pivot', 'Pay Code से grouped summary — EarnAmt, BeneAmt, etc.'],
        ['5. Reconciliation', 'Main sheet — GL vs PR comparison, Variance column, Match/Variance status'],
        ['6. Payroll Process', 'Config mapping rows — GL Code → Pay Code → Amount Column'],
    ]
)

add_para(
    'openpyxl library use होती है। हर sheet में formatting होती है: '
    'colored headers, freeze panes, column widths auto-fit।'
)
doc.add_paragraph()

page_break()

# ════════════════════════════════════════════════════════════════════════════
#  SECTION 9 — MONGODB
# ════════════════════════════════════════════════════════════════════════════
add_heading('9. MongoDB — History कहाँ Save होती है?', 1, C_NAVY)

add_para(
    'MongoDB एक document database है — data JSON-like format में store होता है। '
    'यह tool दो collections use करता है।'
)

add_heading('Collection 1: mapping_configs', 2, C_BLUE)

add_code_box([
    '# mapping_configs collection — एक document:',
    '{',
    '  "_id": ObjectId("..."),',
    '  "client_name": "ClientABC",',
    '  "rows": [',
    '    {',
    '      "recon_step": "A. Earning",',
    '      "gl_code": "5000",',
    '      "gl_title": "Salaries",',
    '      "pay_code": "Wages",',
    '      "pay_code_title": "Regular Wages",',
    '      "amount_column": "EarnAmt",',
    '      "code_type": "EARNING"',
    '    },',
    '    ... (80+ rows)',
    '  ],',
    '  "updated_at": ISODate("2026-01-15T10:30:00Z")',
    '}',
], caption='mapping_configs collection')
doc.add_paragraph()

add_heading('Collection 2: recon_history', 2, C_BLUE)

add_code_box([
    '# recon_history collection — एक run का document:',
    '{',
    '  "_id": ObjectId("..."),',
    '  "client_name": "ClientABC",',
    '  "period_label": "January 2026",',
    '  "gl_filename": "GL_Jan2026.xlsx",',
    '  "pr_filename": "PR_Jan2026.xlsx",',
    '  "gl_row_count": 150,',
    '  "pr_row_count": 2500,',
    '  "summary_stats": {',
    '    "total_lines": 45,',
    '    "matched": 43,',
    '    "variances": 2,',
    '    "total_variance": -150.00,',
    '    "is_clean": false',
    '  },',
    '  "result_data": {',
    '    "recon_table": {"columns": [...], "rows": [...]},',
    '    "gl_pivot":    {"columns": [...], "rows": [...]},',
    '    "pr_pivot":    {"columns": [...], "rows": [...]},',
    '    "unmapped_gl": ["9999"],',
    '    "unmapped_pr": [["BonusNew", "EARNING"]]',
    '  },',
    '  "created_at": ISODate("2026-01-15T10:30:00Z")',
    '}',
], caption='recon_history collection — Full document')
doc.add_paragraph()

page_break()

# ════════════════════════════════════════════════════════════════════════════
#  SECTION 10 — CONNECTION MAP
# ════════════════════════════════════════════════════════════════════════════
add_heading('10. सभी Files का आपस में Connection (Connection Map)', 1, C_NAVY)

add_para('नीचे यह दिखाया है कि कौन सी file किसको call करती है:')
doc.add_paragraph()

add_colored_box(
    '🌐 BROWSER (Frontend)\n'
    '    │\n'
    '    ├── index.html  ──loads──►  app.js (core)\n'
    '    │                              │\n'
    '    │                              ├── upload.js   → /api/upload, /api/confirm-mapping\n'
    '    │                              ├── config.js   → /api/mapping-config, /api/gl-codes, /api/pr-codes\n'
    '    │                              ├── results.js  → /api/run, /api/download\n'
    '    │                              └── history.js  → /api/recon-history\n'
    '    │\n'
    '──────── HTTP ────────────────────────────────────────\n'
    '    │\n'
    '🖥  SERVER (Backend)\n'
    '    │\n'
    '    └── main.py (FastAPI App)\n'
    '            │\n'
    '            ├── routes/upload.py\n'
    '            │       └── file_reader.py  →  column_identifier/\n'
    '            │                                   ├── fuzzy_matcher.py\n'
    '            │                                   ├── bedrock_identifier.py (AWS)\n'
    '            │                                   └── mapping_cache.py (files)\n'
    '            │\n'
    '            ├── routes/columns.py\n'
    '            │       └── state.py (RAM)\n'
    '            │\n'
    '            ├── routes/reconcile.py\n'
    '            │       ├── mapping_config.py  →  db.py / files\n'
    '            │       ├── processors/gl_processor.py\n'
    '            │       ├── processors/payroll_processor.py\n'
    '            │       ├── processors/reconciliation_processor.py\n'
    '            │       ├── utils/excel_exporter.py\n'
    '            │       └── db.py → MongoDB (recon_history)\n'
    '            │\n'
    '            └── routes/mapping_config.py\n'
    '                    └── db.py → MongoDB (mapping_configs)',
    'F0F3F8', C_NAVY
)
doc.add_paragraph()

page_break()

# ════════════════════════════════════════════════════════════════════════════
#  SECTION 11 — DEAD CODE
# ════════════════════════════════════════════════════════════════════════════
add_heading('11. Dead Code — वो Code जो काम नहीं आ रहा', 1, C_NAVY)

add_para(
    '"Dead Code" वो code है जो project में तो है लेकिन कहीं से call या use नहीं होता। '
    'यह project में कुछ ऐसा code है:'
)
doc.add_paragraph()

add_heading('🗂  पूरा Old Streamlit Frontend — बेकार', 2, C_RED)

add_para(
    'यह tool पहले Streamlit (Python web framework) में था। '
    'अब नया HTML/JS frontend बन गया है, लेकिन पुरानी files अभी भी हैं।'
)

add_two_col_table(
    ['Dead File', 'क्यों बेकार है'],
    [
        ['frontend/app.py', 'Old Streamlit app — run.py इसे कभी call नहीं करता'],
        ['frontend/components/file_upload.py', 'Old Streamlit component'],
        ['frontend/components/column_mapping_ui.py', 'Old Streamlit component'],
        ['frontend/components/report_viewer.py', 'Old Streamlit component'],
    ]
)

add_heading('🐍  Python Dead Code', 2, C_RED)

add_two_col_table(
    ['Location', 'Dead Code', 'क्यों बेकार'],
    [
        ['column_identifier/__init__.py:141', 'delete_cached_mapping()', 'सिर्फ old app.py call करती थी'],
        ['routes/upload.py', 'GET /api/sheets endpoint', 'Frontend कभी call नहीं करता'],
        ['config/default_mapping.py:84', 'FIXED_COLUMNS = [...]', 'कहीं import नहीं होता'],
        ['config/default_mapping.py:87', 'EDITABLE_COLUMNS = [...]', 'कहीं import नहीं होता'],
        ['config/default_mapping.py:90', 'AMOUNT_COLUMN_OPTIONS', 'Duplicate — app.js में already है'],
        ['config/default_mapping.py:93', 'CODE_TYPE_OPTIONS', 'Same duplicate'],
    ]
)

add_heading('🌐  HTML Dead Elements', 2, C_RED)

add_two_col_table(
    ['Element', 'क्यों बेकार'],
    [
        ['<input id="cfg-bedrock">', 'upload.js hardcode "true" use करता है, यह input कभी read नहीं होता'],
        ['<input id="cfg-cache">', 'Same — hardcoded है upload.js में'],
    ]
)

add_heading('🎨  CSS Dead Classes', 2, C_RED)

add_two_col_table(
    ['CSS Class', 'क्यों बेकार'],
    [
        ['.pf-warning', 'Old pre-flight CSS — नया .pf-warn use होता है'],
        ['.pf-icon', 'Old CSS — .pf-icon-wrap से replace हो गया'],
        ['.row-alt', 'Result table CSS — JS कभी generate नहीं करता'],
        ['.card-header-navy', 'Upload cards अब .upload-card-header use करते हैं'],
        ['.btn-primary', 'कोई button इस class को use नहीं करता'],
        ['.btn-gold, .btn-white, .btn-danger', 'कहीं use नहीं'],
        ['.badge-green/amber/red/navy/grey', 'Badge classes — कहीं HTML या JS में नहीं'],
        ['.combined-pivot-wrap, .pivot-half', '"Legacy" लिखा है comment में — old layout'],
    ]
)

page_break()

# ════════════════════════════════════════════════════════════════════════════
#  SECTION 12 — SUMMARY IN SIMPLE HINDI
# ════════════════════════════════════════════════════════════════════════════
add_heading('12. सरल भाषा में पूरी प्रक्रिया (Final Summary)', 1, C_NAVY)

add_para(
    'अगर आप सब कुछ भूल गये हों तो बस यह section याद रखें — '
    'इसमें पूरा tool एकदम सरल भाषा में है।'
)
doc.add_paragraph()

add_colored_box(
    '📖 पूरी कहानी — एक बार में:\n\n'
    '1. User browser खोलता है → FastAPI server index.html दिखाता है\n\n'
    '2. Browser /api/session call करता है → Server "abc-123" ID देता है\n\n'
    '3. User GL Report upload करता है:\n'
    '   • File server को जाती है\n'
    '   • Server Excel पढ़ता है (pandas)\n'
    '   • AI (Fuzzy + Bedrock) columns पहचानता है\n'
    '   • Browser को preview और detected columns मिलते हैं\n'
    '   • User confirm करता है → mapping RAM में save\n\n'
    '4. Same process Payroll Register के लिए\n\n'
    '5. Configuration page पर auto-populate होता है:\n'
    '   • GL Codes और Pay Codes file से आते हैं\n'
    '   • GL Code 5000 → "A. Earning" section में auto-place\n'
    '   • User manual changes कर सकता है\n'
    '   • Save करो → MongoDB में जाता है\n\n'
    '6. Run Reconciliation:\n'
    '   • GL और PR दोनों process होते हैं\n'
    '   • Compare होते हैं → Variance निकलता है\n'
    '   • MongoDB history में save\n'
    '   • Browser को results JSON में मिलते हैं\n\n'
    '7. Results page: Green = Match, Red = Variance, Blue = Total\n\n'
    '8. Download: 6-sheet Excel report\n\n'
    '9. नई process: "New Reconciliation" button → session reset → Upload से शुरू',
    'EEF2FF', C_NAVY
)

doc.add_paragraph()

add_heading('Key Technologies जो use होती हैं:', 2, C_BLUE)

add_two_col_table(
    ['Technology', 'काम', 'कहाँ'],
    [
        ['FastAPI (Python)', 'Web server और API endpoints', 'backend/api/'],
        ['pandas (Python)', 'Excel/CSV पढ़ना और process करना', 'processors/'],
        ['openpyxl (Python)', 'Excel export बनाना', 'utils/excel_exporter.py'],
        ['rapidfuzz (Python)', 'String similarity matching', 'column_identifier/'],
        ['boto3 (Python)', 'AWS Bedrock AI call करना', 'column_identifier/'],
        ['pymongo (Python)', 'MongoDB connection', 'backend/api/db.py'],
        ['HTML + CSS', 'User interface बनाना', 'frontend/'],
        ['JavaScript', 'Browser में logic चलाना', 'frontend/static/js/'],
        ['MongoDB', 'History और config persistent store', 'Database'],
    ]
)

doc.add_paragraph()

add_colored_box(
    '💬 अंत में एक बात:\n'
    'यह tool एकदम modular है — हर चीज़ अपनी जगह है:\n'
    '• File reading बदलनी हो → file_reader.py\n'
    '• Column detection बदलनी हो → column_identifier/\n'
    '• GL processing बदलनी हो → gl_processor.py\n'
    '• UI बदलनी हो → frontend/ files\n'
    '• Database बदलनी हो → db.py\n\n'
    'एक जगह change करने से बाकी जगह कुछ नहीं टूटता — यही good design है।',
    'ECFDF5', C_GREEN
)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('— दस्तावेज़ समाप्त —')
r.bold = True
r.font.size = Pt(12)
r.font.color.rgb = C_GREY
r.font.name = 'Mangal'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f'Payroll Reconciliation Tool  ·  Code Documentation  ·  {datetime.date.today().strftime("%d %B %Y")}')
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0xC8, 0xD3, 0xE3)

# ── Save ────────────────────────────────────────────────────────────────────
output_path = r"c:\My Work\payroll reconcilation\payroll_Reconciliation\Payroll_Recon_Code_Documentation_Hindi.docx"
doc.save(output_path)
print(f"✅ Document saved: {output_path}")
